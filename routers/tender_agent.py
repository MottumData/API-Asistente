""" En este modulo se define la funcionalidad para el agente de licitaciones. """

import os
import asyncio
import logging
import time
from typing import List, Dict
from pydantic import BaseModel

from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import PromptTemplate
from langchain_openai import AzureChatOpenAI, AzureOpenAIEmbeddings
from openai import RateLimitError
from fastapi import APIRouter, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse
from langchain_community.document_loaders import PyPDFLoader
from langchain_chroma import Chroma
from docx import Document
import markdown2
from bs4 import BeautifulSoup


from ..internal.path_utils import TEMP_DIR, CHROMA_DB, CHROMA_COLLECTION
from ..internal.prompt_utils import load_prompt
from .rag_utils import load_document_chroma

router = APIRouter()

logger = logging.getLogger(__name__)

llm = AzureChatOpenAI(
    azure_deployment='gpt-4o',
    api_version="2024-08-01-preview",
    azure_endpoint=os.getenv("AZURE_ENDPOINT"),
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    temperature=0.2  # Ajusta según tus necesidades

)

embeddings = AzureOpenAIEmbeddings(
    model="LLM-Codexca_text-embedding-3-large",
    dimensions=None,
    api_key=os.getenv("AZURE_API_KEY"),
    azure_endpoint=os.getenv("AZURE_ENDPOINT"),
    api_version="2023-05-15"
)


class RelevantDocumentRequest(BaseModel):
    """Modelo para la solicitud de documentos relevantes."""
    query: str
    num_proposals: int = 5


class ConceptNotesRequest(BaseModel):
    """Modelo para la solicitud de concept notes."""
    proposal_id: str
    tdr_summary: str
    information_sources: List[str]


class SaveConceptNotesRequest(BaseModel):
    """Modelo para la solicitud de concept notes."""
    proposal_id: str
    concept_notes: Dict


class IndexRequest(BaseModel):
    """Modelo para la solicitud del índice de la propuesta"""
    proposal_id: str


class SaveIndexRequest(BaseModel):
    """Modelo para la solicitud de guardado del índice de la propuesta"""
    proposal_id: str
    index: Dict


class ContentRequest(BaseModel):
    """Modelo para la solicitud de generación de contenido de la propuesta"""
    proposal_id: str


class SaveContentRequest(BaseModel):
    """Modelo para la solicitud de guardado del contenido de la propuesta"""
    proposal_id: str
    content: Dict


class DownloadProposalRequest(BaseModel):
    """Modelo para la solicitud de descarga de la propuesta"""
    proposal_id: str


class SetTitleRequest(BaseModel):
    """Modelo para la solicitud de guardado del contenido de la propuesta"""
    proposal_id: str
    title: str


class TenderProposal(BaseModel):
    """Modelo para las propuestas de proyecto."""
    proposal_id: str
    tdr_summary: str = None
    title: str = None
    information_sources: List[str] = None
    concept_notes: Dict[str, str] = None
    index: Dict = None
    key_ideas: str = None
    content: Dict[str, str] = None

    def get_proposal_id(self):
        return self.proposal_id

    def set_proposal_id(self, proposal_id: str):
        self.proposal_id = proposal_id

    def get_tdr_summary(self):
        return self.tdr_summary

    def set_tdr_summary(self, tdr_summary: str):
        self.tdr_summary = tdr_summary

    def get_title(self):
        return self.title

    def set_title(self, title: str):
        self.title = title

    def get_information_sources(self):
        return self.information_sources

    def set_information_sources(self, information_sources: List[str]):
        self.information_sources = information_sources

    def get_concept_notes(self):
        return self.concept_notes

    def set_concept_notes(self, concept_notes: str):
        self.concept_notes = concept_notes

    def get_index(self):
        return self.index

    def set_index(self, index: List[str]):
        self.index = index

    # def get_key_ideas(self):
    #     return self.key_ideas

    # def set_key_ideas(self, key_ideas: str):
    #     self.key_ideas = key_ideas

    def get_content(self):
        return self.content

    def set_content(self, content: Dict[str, str]):
        self.content = content


proposals: Dict[str, TenderProposal] = {}


def safe_check(doc: dict, key: str):
    """ Verifica si una clave está presente en un diccionario. """
    if key in doc.keys():
        return True
    return False


def get_tender_data(doc):
    """ Obtiene los datos especificados en el esquema de las licitaciones. """
    schema = load_prompt(prompt_name="tender_schema")
    json_template = load_prompt(prompt_name="extract_tender_data")
    parser = JsonOutputParser()

    prompt = PromptTemplate(
        template=(json_template),
        input_variables=['schema', 'format_instructions', 'document'],
        partial_variables={
            "format_instructions": parser.get_format_instructions(),
        }
    )
    chain = prompt | llm | parser
    tdr_data = chain.invoke({"document": doc, "schema": schema})
    return tdr_data


def make_summary(doc, tender_data):
    """ Genera un resumen  a partir de los términos de referencia y la información extraída. """
    json_template = load_prompt(prompt_name="make_tender_summary")
    summary_prompt = PromptTemplate(
        template=(json_template),
        input_variables=['document', 'tender_data'],
    )
    chain = summary_prompt | llm
    summary = chain.invoke({"document": doc, "tender_data": tender_data})
    return summary


@router.post("/upload-tdr/")
async def upload_tdr(file: UploadFile = File(...), num_proposals: int = 5):
    """ Sube un documento y lo guarda temporalmente en el directorio temp. """

    if not os.path.exists(TEMP_DIR):
        os.makedirs(TEMP_DIR)
    temp_file_path = os.path.join(TEMP_DIR, file.filename)

    with open(temp_file_path, "wb") as temp_file:
        contents = await file.read()
        temp_file.write(contents)

    loader = PyPDFLoader(temp_file_path)
    document = loader.load()

    tender_data = get_tender_data(document)
    summary = make_summary(document, tender_data)
    related_project = make_project_proposal(
        RelevantDocumentRequest(query=summary.content,
                                num_proposals=num_proposals),
        from_endpoint=False)

    os.remove(temp_file_path)

    return JSONResponse({"key points": tender_data, "complete summary": summary.content,
                         "related projects": related_project})


@router.post('/related-projects/')
def make_project_proposal(request: RelevantDocumentRequest, from_endpoint: bool = False):
    """ Obtiene hasta X propuestas de proyecto más afines a partir de los términos de referencia
        y la información extraída. """

    query = request.query
    num_proposals = request.num_proposals
    vector_store = Chroma(embedding_function=embeddings,
                          persist_directory=CHROMA_DB,
                          collection_name=CHROMA_COLLECTION)

    retriever = vector_store.as_retriever(
        search_type='mmr', search_kwargs={'k': num_proposals})

    related_docs = retriever.invoke(query)

    # Devuelve los documentos únicos
    unique_docs = {}
    for doc in related_docs:
        source = str(doc.metadata['source']).split(sep='\\')[-1]
        if source not in unique_docs:
            unique_docs[source] = doc.page_content
        else:
            unique_docs[source] += doc.page_content

    related_docs = [{source: content}
                    for source, content in unique_docs.items()]
    if from_endpoint:
        return JSONResponse({"related_projects": related_docs, "query": query})
    return related_docs

# TODO - Rediseñar la forma en que se cargan los documentos de información.

@router.post('/make-concept-note/')
def make_concept_notes(request: ConceptNotesRequest):
    """ Genera las notas conceptuales de un proyecto. """

    proposal_id = str(request.proposal_id)
    information_sources = request.information_sources
    tdr_summary = request.tdr_summary

    logger.info('Creating proposal with ID: %s', proposal_id)

    proposal = TenderProposal(proposal_id=proposal_id)
    proposal.set_information_sources(information_sources)
    proposal.set_tdr_summary(tdr_summary)
    proposals[proposal_id] = proposal  # Crea una nueva propuesta

    while not safe_check(proposals, proposal_id):
        time.sleep(0.2)

    logger.info('Proposal created with ID: %s', proposal_id)

    logger.info('Generating concept notes for proposal: %s. Information sources: %s',
                proposal_id, str(information_sources))
    information_source_docs = [load_document_chroma(source)
                               for source in information_sources]

    json_template = load_prompt(prompt_name="make_concept_note")
    json_parser = JsonOutputParser()
    concept_notes_prompt = PromptTemplate(
        template=(json_template),
        input_variables=['tdr_summary', 'information_sources'],
        partial_variables={
            "format_instructions": json_parser.get_format_instructions(),
        }
    )
    chain = concept_notes_prompt | llm | json_parser
    try:
        response = chain.invoke(
            {"tdr_summary": tdr_summary, "information_sources": information_source_docs})
    except RateLimitError as e:
        raise HTTPException(
            status_code=429,
            detail="Documento demasiado extenso para el modelo. \
                Pruebe con un documento más corto.") from e
    return JSONResponse({"concept_notes": response})


@router.post('/save-concept-note/')
async def save_concept_note(request: SaveConceptNotesRequest):
    """ Guarda las notas conceptuales de un proyecto a partir de su ID. """

    proposal_id = str(request.proposal_id)
    concept_notes = request.concept_notes

    proposal = proposals[proposal_id]
    proposal.set_concept_notes(concept_notes)
    proposals[proposal_id] = proposal
    return JSONResponse(status_code=200,
                        content={"message": "Notas conceptuales guardadas con éxito."})


@router.post('/make-index/')
async def make_index(request: IndexRequest):
    """ Genera un índice de un proyecto. """
    proposal_id = str(request.proposal_id)
    proposal = proposals[proposal_id]
    concept_notes = proposal.get_concept_notes()
    tdr_summary = proposal.get_tdr_summary()

    json_template = load_prompt(prompt_name="make_index")
    json_parser = JsonOutputParser()

    index_prompt = PromptTemplate(
        template=(json_template),
        input_variables=['tdr_summary', 'concept_notes'],
        partial_variables={
            "format_instructions": json_parser.get_format_instructions(),
        }
    )

    chain = index_prompt | llm | json_parser
    response = chain.invoke(
        {"tdr_summary": tdr_summary, "concept_notes": concept_notes})
    return response


@router.post('/save-index/')
async def save_index(request: SaveIndexRequest):
    """ Guarda el índice de un proyecto a partir de su ID. """
    proposal_id = str(request.proposal_id)
    index = request.index
    proposal = proposals[proposal_id]
    proposal.set_index(index)
    proposals[proposal_id] = proposal
    return JSONResponse(status_code=200, content={"message": "Índice guardado con éxito."})


@router.post('/make-content/')
async def make_proposal_content(request: ContentRequest):
    """ Genera el contenido de una propuesta a partir de su ID. """
    proposal = proposals[str(request.proposal_id)]
    concept_notes = proposal.get_concept_notes()
    tdr_summary = proposal.get_tdr_summary()
    index = proposal.get_index()
    business_information = load_document_chroma(
        'CODEXCA - DOSSIER DIGITAL 2021.pdf')
    content = {}
    json_template = load_prompt(prompt_name="project_writer")
    i = 1
    for ind, resume in index.items():
        content_prompt = PromptTemplate(
            template=(json_template),
            input_variables=['index', 'description', 'concept_notes',
                             'business_information', 'tdr_summary'],
        )
        chain = content_prompt | llm
        logger.info('Generating content for proposal: %s. Index name: %s - %s/%s',
                    proposal.get_proposal_id(), ind, i, len(index))
        response = chain.invoke(
            {"index": ind, "description": resume, "concept_notes": concept_notes,
             "business_information": business_information, "tdr_summary": tdr_summary})
        content[ind] = response.content
        i += 1
    return JSONResponse({"content": content})


@router.post('/save-content/')
async def save_proposal_content(request: SaveContentRequest):
    """ Guarda el contenido de una propuesta a partir de su ID. """
    proposal_id = str(request.proposal_id)
    content = request.content
    proposal = proposals[proposal_id]
    proposal.set_content(content)
    proposals[proposal_id] = proposal
    return JSONResponse(status_code=200, content={"message": "Contenido guardado con éxito."})


@router.get('/proposal-trace/')
async def get_proposal_trace(proposal_id: str):
    """ Obtiene el registro de una propuesta a partir de su ID."""
    try:
        proposal = proposals[proposal_id]
    except KeyError:
        return JSONResponse(status_code=404, content={"message": "Código de propuesta no encontrado."})
    return JSONResponse(content=proposal.model_dump(mode='json'))


@router.post('/set-title/')
async def set_proposal_title(request: SetTitleRequest):
    """ Guarda el título de una propuesta a partir de su ID. """
    proposal_id = request.proposal_id
    title = request.title
    proposal = proposals[proposal_id]
    proposal.set_title(title)
    proposals[proposal_id] = proposal
    return JSONResponse(status_code=200, content={"message": "Título guardado con éxito."})


@router.post('/download-proposal/')
async def download_proposal(request: DownloadProposalRequest, background_tasks: BackgroundTasks):
    """ Descarga una propuesta en .docx a partir de su ID. """
    proposal_id = request.proposal_id
    proposal = proposals[proposal_id]
    title = proposal.get_title() if proposal.get_title() else proposal_id
    content = proposal.get_content()
    doc = Document()
    # for _, text in content.items():
    #     html = markdown2.markdown(text)
    #     temp_html_path = os.path.join(TEMP_DIR, 'temp.html')
    #     with open(temp_html_path, 'w', encoding='utf-8') as temp_html_file:
    #         temp_html_file.write(html)
    #     temp_docx_path = os.path.join(TEMP_DIR, 'temp.docx')
    #     pypandoc.convert_file(
    #         temp_html_path, 'docx', format='html', outputfile=temp_docx_path)
    #     docx_content = Document(temp_docx_path)
    #     for paragraph in docx_content.paragraphs:
    #         doc.add_paragraph(paragraph.text)
    for _, text in content.items():
        html = markdown2.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')
        for element in soup:
            if element.name == 'p':
                paragraph = doc.add_paragraph()
                for part in element:
                    if part.name == 'strong':
                        run = paragraph.add_run(part.text)
                        run.bold = True
                    else:
                        paragraph.add_run(part.text)
    path = os.path.join(TEMP_DIR, f'{title}.docx')
    doc.save(path)
    background_tasks.add_task(os.remove, path)
    return FileResponse(path, filename=f'{title}.docx',
                        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@router.post("/upload-tdr-streaming/")
async def upload_tdr_streaming(file: UploadFile = File(...)):
    """ Sube un documento y lo guarda temporalmente en el directorio temp. """

    if not os.path.exists(TEMP_DIR):
        os.makedirs(TEMP_DIR)
    temp_file_path = os.path.join(TEMP_DIR, file.filename)

    with open(temp_file_path, "wb") as temp_file:
        contents = await file.read()
        temp_file.write(contents)

    loader = PyPDFLoader(temp_file_path)
    document = loader.load()

    async def generate():
        tender_data = get_tender_data(document)
        yield JSONResponse({"key points": tender_data}).body.decode()
        await asyncio.sleep(1)

        summary = make_summary(document, tender_data)
        yield JSONResponse({"complete summary": summary.content}).body.decode()
        await asyncio.sleep(1)

        # Elimina el archivo temporal después de procesarlo
        os.remove(temp_file_path)

    return StreamingResponse(generate(), media_type="application/json")

# TODO - Implementar logica para proposal_id generalizado
